import os
import zipfile
import sys
import docx
import openpyxl
import xlrd
import subprocess



def readXls(path):
    file_name = path
    book = xlrd.open_workbook(file_name)
    sheetNames = book.sheet_names()
    res = ''
    for sheetName in sheetNames:
        sheet = book.sheet_by_name(sheetName)
        rowCounts = sheet.nrows
        for i in range(rowCounts):
            data = sheet.row(i)
            for cell in data:
                if cell.value != '':
                    res = res + str(cell.value)
    res = res.replace('\n', '')
    print(res)


def readXlsx(path):
    file_name = path
    book = openpyxl.load_workbook(file_name)
    res = ''
    for sheet in book:
        for row in sheet.rows:
            for cell in row:
                if cell.value is not None:
                    res = res + str(cell.value)
    res = res.replace('\n', '')
    print(res)


def readDocx(path):
    file_name = path
    # file_name='训练行为监督综述初稿.docx'
    file = docx.Document(file_name)
    res = ''
    # 输出每一段的内容
    for para in file.paragraphs:
        if para.text != '':
            res = res + para.text
    print(res)

    # 表格内容
    tableRes = ''
    for table in file.tables:
        for row in table.rows:  # 读每行
            for cell in row.cells:  # 读一行中的所有单元格
                if cell.text != '':
                    tableRes = tableRes + cell.text
    print(tableRes)


# 利用antiword 只在Linux环境可用
def readDoc(path):
    word = path
    output = subprocess.check_output(["antiword", word])
    print(output)


def chengeChar(path):
    '''处理乱码'''
    if not os.path.exists(path): return path
    path = path.rstrip('/').rstrip('\\')  # 去除路径最右边的/
    file_name = os.path.split(path)[-1]  # 获取最后一段字符，准备转换
    file_path = os.path.split(path)[0]  # 获取前面的路径，为rename做准备
    try:  # 将最后一段有乱码的字符串转换，尝试过整个路径转换，不生效，估计是无法获取整个路径的编码格式吧。
        new_name = file_name.encode('cp437').decode('gbk')
    except:  # 先转换成Unicode再转换回gbk或utf-8
        new_name = file_name.encode('utf-8').decode('utf-8')
    path2 = os.path.join(file_path, new_name)  # 将转换完成的字符串组合成新的路径
    try:
        os.renames(path, path2)  # 重命名文件
    except:
        print('renames error！！')
    return path2


def del_zip(path):
    '''删除解压出来的zip包'''
    path = chengeChar(path)
    if path.endswith('.zip') or path.endswith('.rar'):
        os.remove(path)
    elif os.path.isdir(path):
        for i in os.listdir(path):
            file_path = os.path.join(path, i)
            del_zip(file_path)  # 递归调用，先把所有的文件删除


def unzip_file(z, unzip_path):
    '''解压zip包'''
    z.extractall(path=unzip_path)
    zip_list = z.namelist()  # 返回解压后的所有文件夹和文件list
    z.close()
    for zip_file in zip_list:
        path = os.path.join(unzip_path, zip_file)
        if os.path.exists(path): readZip(path)


def readZip(path):
    '''主逻辑函数'''
    list = []
    path = chengeChar(path)
    # print(path)
    if os.path.exists(path):
        unzip_path = os.path.split(path)[0]  # 解压至当前目录
        list.append(path)
        if path.endswith('.zip'):
            z = zipfile.ZipFile(path, 'r')
            unzip_file(z, unzip_path)
        elif path.endswith('.rar'):
            r = rarfile.RarFile(path)
            unzip_file(r, unzip_path)
        elif os.path.isdir(path):
            for file_name in os.listdir(path):
                path = os.path.join(path, file_name)
                if os.path.exists(path):
                    readZip(path)
                    os.remove(path)
        else:
            # print(path)
            if path.endswith('.doc'):
                print(path)
                # readDoc(path)
            elif path.endswith('.xls'):
                readXls(path)
            elif path.endswith('.xlsx'):
                readXlsx(path)
            elif path.endswith('.docx'):
                readDocx(path)

    else:
        print('the path is not exist!!!')
    for i in list:
        os.remove(i)


if __name__ == '__main__':
    # zip_path = r'C:\Users\ES-IT-PC-193\Desktop\aa\A.zip'
    # zip_path = 'trash/zip/zip.zip'  # 接收传入的路径参数
    # readZip(zip_path)
    # if zipfile.is_zipfile(zip_path):  # 删除解压出来的压缩包
    #     del_zip(os.path.splitext(zip_path)[0])  # 以后缀名切割
    # readXls()
    readDocx('trash/data.docx')
